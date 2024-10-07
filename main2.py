import os
from copy import deepcopy

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    產生 document 或 table cell 中的所有段落和表格
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("未知的父級物件")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def copy_block(block, new_doc):
    """
    將段落或表格直接複製到新的文檔中，保留格式
    """
    if isinstance(block, Paragraph):
        new_paragraph = deepcopy(block._p)
        new_doc.element.body.append(new_paragraph)
    elif isinstance(block, Table):
        new_table = deepcopy(block._tbl)
        new_doc.element.body.append(new_table)

def split_docx_by_keyword(input_file, keyword):
    # 檢查檔案是否存在
    if not os.path.exists(input_file):
        print(f"檔案 {input_file} 不存在，請確認檔案名稱和路徑是否正確。")
        return []

    # 讀取原始 Word 文件
    doc = Document(input_file)
    output_files = []
    new_doc = Document()
    section_index = 1  # 用於命名輸出文件
    found_keyword = False  # 用於檢查是否遇到關鍵字

    # 清除 new_doc 的默認內容
    new_doc._element.remove_all()

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            # 檢查段落中是否包含關鍵字
            if keyword in block.text:
                # 如果 new_doc 不是空的，先保存之前的內容
                if found_keyword or len(new_doc.element.body):
                    output_filename = f"section_{section_index}.docx"
                    new_doc.save(output_filename)
                    output_files.append(output_filename)
                    section_index += 1
                    # 創建新的 Document 物件
                    new_doc = Document()
                    new_doc._element.remove_all()
                found_keyword = True  # 標記已經找到關鍵字

            # 將段落直接複製到新的文檔中
            copy_block(block, new_doc)

        elif isinstance(block, Table):
            # 檢查表格中是否包含關鍵字
            table_text = ''
            for row in block.rows:
                for cell in row.cells:
                    table_text += cell.text
            if keyword in table_text:
                if found_keyword or len(new_doc.element.body):
                    output_filename = f"section_{section_index}.docx"
                    new_doc.save(output_filename)
                    output_files.append(output_filename)
                    section_index += 1
                    new_doc = Document()
                    new_doc._element.remove_all()
                found_keyword = True

            # 將表格直接複製到新的文檔中
            copy_block(block, new_doc)

    # 保存最後一部分內容
    if len(new_doc.element.body):
        output_filename = f"section_{section_index}.docx"
        new_doc.save(output_filename)
        output_files.append(output_filename)

    return output_files


# 使用示例
if __name__ == "__main__":
    # 讓使用者輸入檔案名稱和關鍵字
    input_file = input("請輸入要處理的 Word 檔案名稱（包含路徑）：")
    keyword = input("請輸入要查找的關鍵字：")

    # 呼叫函數並顯示結果
    split_files = split_docx_by_keyword(input_file, keyword)
    if split_files:
        print("已生成以下分割檔案：")
        for filename in split_files:
            print(filename)
    else:
        print("未找到關鍵字，或未生成任何檔案。")