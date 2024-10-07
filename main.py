from copy import deepcopy

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def split_word_pages(input_file):
    # 開啟原始 Word 文件
    doc = Document(input_file)

    # 計數器，用於命名每個分頁檔案
    page_number = 1

    # 建立新文件來存放每一頁的內容
    new_doc = Document()
    for paragraph in doc.paragraphs:
        new_doc.add_paragraph(paragraph.text)

        # 判斷是否該換頁 (在此例中，假設以空段落分隔頁面)
        if not paragraph.text.strip():
            output_file = f"page_{page_number}.docx"
            new_doc.save(output_file)
            print(f"已儲存 {output_file}")
            page_number += 1
            new_doc = Document()  # 建立新頁文件

    # 儲存最後一頁
    if new_doc.paragraphs:
        output_file = f"page_{page_number}.docx"
        new_doc.save(output_file)
        print(f"已儲存 {output_file}")

def iter_block_items(parent):
    """
    生成 document 或 table cell 中的所有段落、表格和圖片。
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("未知的父級物件")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def copy_run_format(source_run, target_run):
    """
    複製 Run 的格式。
    """
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.style = source_run.style
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb

def copy_paragraph(paragraph, new_doc):
    """
    將段落複製到新的文檔中。
    """
    new_para = new_doc.add_paragraph()
    new_para.style = paragraph.style
    for run in paragraph.runs:
        new_run = new_para.add_run(run.text)
        copy_run_format(run, new_run)
    # 複製圖片
    for inline_shape in paragraph.runs:
        if 'graphic' in inline_shape._r.xml:
            new_run = new_para.add_run()
            new_run._r.append(deepcopy(inline_shape._r))
            copy_run_format(inline_shape, new_run)

def copy_table(table, new_doc):
    """
    將表格複製到新的文檔中。
    """
    tbl = new_doc.add_table(rows=0, cols=len(table.columns))
    tbl.style = table.style
    for row in table.rows:
        new_row = tbl.add_row()
        for idx, cell in enumerate(row.cells):
            new_cell = new_row.cells[idx]
            copy_cell_contents(cell, new_cell)

def copy_cell_contents(source_cell, target_cell):
    """
    複製表格單元格的內容。
    """
    for paragraph in source_cell.paragraphs:
        new_para = target_cell.add_paragraph()
        new_para.style = paragraph.style
        for run in paragraph.runs:
            new_run = new_para.add_run(run.text)
            copy_run_format(run, new_run)
    for table in source_cell.tables:
        copy_table(table, target_cell)

def split_docx_by_keyword(input_file, keyword):
    # 讀取原始 Word 文件
    doc = Document(input_file)
    output_files = []
    new_doc = Document()
    section_index = 1  # 用於命名輸出文件
    found_keyword = False  # 用於檢查是否遇到關鍵字

    # 清除 new_doc 的默認內容
    new_doc._element.remove_all()

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            # 檢查段落中是否包含關鍵字
            if keyword in block.text:
                # 如果 new_doc 不是空的，先保存之前的內容
                if found_keyword or len(new_doc.element.body):
                    output_filename = f"section_{section_index}.docx"
                    new_doc.save(output_filename)
                    output_files.append(output_filename)
                    section_index += 1
                    # 創建新的 Document 物件
                    new_doc = Document()
                    new_doc._element.clear_content()
                found_keyword = True  # 標記已經找到關鍵字

            # 將段落添加到新的文檔中
            copy_paragraph(block, new_doc)

        elif isinstance(block, Table):
            # 檢查表格中是否包含關鍵字
            table_text = ''
            for row in block.rows:
                for cell in row.cells:
                    table_text += cell.text
            if keyword in table_text:
                if found_keyword or len(new_doc.element.body):
                    output_filename = f"section_{section_index}.docx"
                    new_doc.save(output_filename)
                    output_files.append(output_filename)
                    section_index += 1
                    new_doc = Document()
                    new_doc._element.remove_all()
                found_keyword = True

            # 將表格添加到新的文檔中
            copy_table(block, new_doc)

    # 保存最後一部分內容
    if len(new_doc.element.body):
        output_filename = f"section_{section_index}.docx"
        new_doc.save(output_filename)
        output_files.append(output_filename)

    return output_files


# 使用示例
if __name__ == "__main__":
    # split_word_pages("000.docx")
    split_docx_by_keyword("000.docx", "預  約  單")